import os
from django.db import models
from django.core.exceptions import ValidationError
from django.utils.translation import gettext_lazy as _
from django.utils.html import format_html, format_html_join
from docx import Document
import zipfile
import Levenshtein
from .wordChoices import *
from .fileModels import *
from pathlib import Path

BASE_DIR = Path(__file__).resolve().parent.parent
MEDIA_ROOT  = BASE_DIR / 'media'


def getColororNone(c):
    if c:
        return c
    return c.rgb

def getPtorNone(sz):
    if sz:
        return (round(sz.pt))
    return 0

def getSzorNone(sz1, sz2):
    if sz1:
        return (round(sz1.pt))
    if sz2:
        return (round(sz2.pt))
    return 0

def getLineorNone(sz1, sz2):
    if sz1:
        return (round(sz1))
    if sz2:
        return (round(sz2))
    return 0

# Create your models here.
class WordQuestion(models.Model):

    def __str__(self):
        return 'Word操作'+str(self.id)

    class Meta:
        verbose_name = '题目-Word操作'
        verbose_name_plural = '题目-Word操作'

    def file_path(self):
        return self.upload_docx.name
    file_path.short_description = 'Word文件'

    def base_path_(self):
        return os.path.join(MEDIA_ROOT, 'upload_docx',str(self.id))

    def zip_path_(self):
        return os.path.join(MEDIA_ROOT, 'upload_docx',str(self.id),'word操作题目.zip')

    def save(self, *args, **kwargs):
        super().save(*args, **kwargs)  
        if not os.path.exists(self.base_path_()):  os.mkdir(self.base_path_())

        zip_path = self.zip_path_()
        if zip_path:
            if os.path.exists(zip_path): os.remove(zip_path)
            zf = zipfile.ZipFile(zip_path, 'w')
            zf.write(self.upload_docx.path,'{}/word.docx'.format('word操作题目'))
            if self.image_op:
                zf.write(self.upload_image_file.path,'{}/{}'.format('word操作题目', self.upload_image_file.name.split('/')[-1]))
            zf.close()

    def word_op_numb(self):
        op_list = []
        if self.char_edit_op:
            op_list.append('查找替换')
        if self.font_op:
            op_list.append('字体设置')
        if self.paraformat_op:
            op_list.append('段落设置')
        if self.style_op:
            op_list.append('样式设置')
        if self.image_op:
            op_list.append('图片插入')
        if self.table_op:
            op_list.append('表格插入')
        return str(len(op_list))
    word_op_numb.short_description = '操作数目'

    def para_text_simple(self, origin_text):
        if len(origin_text)<30:
            return origin_text.strip()
        else:
            return origin_text[:15] +' ...... '+origin_text[-15:]
    para_text_simple.short_description = '考查段落内容'
    
    def operation_description_all(self):
        description_list = [self.char_edit_description(), 
        self.font_description(),
        self.paraformat_description(), 
        self.style_description(),
        self.image_description(),
        self.table_description(),
        ]
        description_list = [x for x in description_list if len(x)>0]
        return format_html("<ol>") + \
                format_html_join(
                '\n', '<li style="color:black;">{}</li>',
                ((x,) for x in description_list)
                ) \
                + format_html("</ol>")
    operation_description_all.short_description = '题目文字描述'

    def char_edit_description(self):
        if self.char_edit_op:
            desc_all = '将段落【'+self.para_text_simple(self.char_edit_text)+'】'\
                +'中所有“'+self.char_edit_origin+'”替换成“'+self.char_edit_replace+'”'+'.'
            return desc_all
        else:
            return ''
    char_edit_description.short_description = '查找替换文字描述'

    def font_description(self):
        if self.font_op:
            setting_list=[]
            if self.font_name_chinese !='': setting_list.append('中文'+self.font_name_chinese)
            if self.font_name_ascii !='':  setting_list.append('西文'+self.font_name_ascii)
            if self.font_size !='': setting_list.append('字号'+self.font_size+'磅')
            if self.font_color !='': setting_list.append('标准色'+[x[1] for x in FONT_COLOR_CHOICES if x[0]==self.font_color][0])
            if self.font_bold==True: setting_list.append('粗体')
            if self.font_italic==True: setting_list.append('斜体')
            if self.font_underline !='': 
                setting_list.append([x[1] for x in FONT_UNDERLINE_CHOICES if x[0]==self.font_underline][0])
            desc_all = '将段落【'+self.para_text_simple(self.font_text)+'】'\
                +'字体设置成'+'、'.join(setting_list)+'.'
            return desc_all
        else:
            return ''
    font_description.short_description = '字体设置描述'

    def paraformat_description(self):
        if self.paraformat_op:
            setting_list=[]
            if self.para_alignment !='': 
                setting_list.append([x[1] for x in PARA_ALIGNMENT_CHOICES if x[0]==self.para_alignment][0])
                # setting_list.append(self.para_alignment)
            if self.para_left_indent !='': setting_list.append('左缩进'+self.para_left_indent+'磅')
            if self.para_right_indent !='': setting_list.append('右缩进'+self.para_right_indent+'磅')
            if self.para_first_line_indent !='' and self.para_first_line_indent_size !='': 
                setting_list.append(self.para_first_line_indent+self.para_first_line_indent_size+'磅')
            if self.para_space_before !='': setting_list.append('段前'+self.para_space_before+'磅')
            if self.para_space_after !='':  setting_list.append('段后'+self.para_space_after+'磅')
            if self.para_line_spacing_rule !='': 
                if self.para_line_spacing_rule in (x[0] for x in PARA_LINE_SPACING_RULE_CHOICES[:3]):
                    setting_list.append([x[1] for x in PARA_LINE_SPACING_RULE_CHOICES if x[0]==self.para_line_spacing_rule][0])
                else:
                    setting_list.append([x[1] for x in PARA_LINE_SPACING_RULE_CHOICES if x[0]==self.para_line_spacing_rule][0]+self.para_line_spacing+'倍行距')
            
            # if self.para_firstchardropcap !='' and self.para_firstchardropcaplines !='': 
            #     if self.para_firstchardropcap == PARA_FIRSTCHARDROPCAP_CHOICES[0][0]:
            #         setting_list.append('首字下沉'+self.para_firstchardropcaplines+'行')
            #     if self.para_firstchardropcap == PARA_FIRSTCHARDROPCAP_CHOICES[1][0]:
            #         setting_list.append('首字悬挂'+self.para_firstchardropcaplines+'行')
                
            if self.page_break_before==True: setting_list.append('段前分页')
            if self.keep_with_next==True: setting_list.append('与下段同页')
            if self.keep_together==True: setting_list.append('段中不分页')
            if self.widow_control==True: setting_list.append('孤行控制')

            desc_all = '将段落【'+self.para_text_simple(self.paraformat_text)+'】'\
                +'其段落格式设置成'+'、'.join(setting_list)+'.'
            return desc_all
        else:
            return ''
    paraformat_description.short_description = '段落设置描述'

    def style_description(self):
        if self.style_op:
            style_des = ''
            style_des_add = []
            if self.style_name in ['新样式1', '新样式2']:
                style_des = '创建并应用“'+self.style_name+'”'
            else:
                style_des = '应用样式“'+self.style_name+'”'

            font_setting_list=[]
            if self.style_font_name_chinese !='': font_setting_list.append('中文'+self.style_font_name_chinese)
            if self.style_font_name_ascii !='':  font_setting_list.append('西文'+self.style_font_name_ascii)
            if self.style_font_size !='': font_setting_list.append('字号'+self.style_font_size+'磅')
            if self.style_font_color !='': font_setting_list.append('标准色'+[x[1] for x in FONT_COLOR_CHOICES if x[0]==self.style_font_color][0])
            if self.style_font_bold==True: font_setting_list.append('粗体')
            if self.style_font_italic==True: font_setting_list.append('斜体')
            if self.style_font_underline !='': 
                font_setting_list.append([x[1] for x in FONT_UNDERLINE_CHOICES if x[0]==self.style_font_underline][0])
            if len(font_setting_list)>0:
                style_des_add.append('其字体设置成'+'、'.join(font_setting_list))

            para_setting_list=[]
            if self.style_para_alignment !='':
                #  para_setting_list.append(self.style_para_alignment)
                 para_setting_list.append([x[1] for x in PARA_ALIGNMENT_CHOICES if x[0]==self.style_para_alignment][0])

            if self.style_para_left_indent !='': para_setting_list.append('左缩进'+self.style_para_left_indent+'磅')
            if self.style_para_right_indent !='': para_setting_list.append('右缩进'+self.style_para_right_indent+'磅')
            if self.style_para_first_line_indent !='' and \
               self.style_para_first_line_indent_size !='': 
                para_setting_list.append(self.style_para_first_line_indent+self.style_para_first_line_indent_size+'磅')
            if self.style_para_space_before !='': para_setting_list.append('段前'+self.style_para_space_before+'磅')
            if self.style_para_space_after !='':  para_setting_list.append('段后'+self.style_para_space_before+'磅')
            if self.style_para_line_spacing_rule !='': 
                if self.style_para_line_spacing_rule in (x[0] for x in PARA_LINE_SPACING_RULE_CHOICES[:3]):
                    para_setting_list.append([x[1] for x in PARA_LINE_SPACING_RULE_CHOICES if x[0]==self.style_para_line_spacing_rule][0])
                else:
                    para_setting_list.append([x[1] for x in PARA_LINE_SPACING_RULE_CHOICES if x[0]==self.style_para_line_spacing_rule][0]+self.para_line_spacing+'倍行距')
                # if self.style_para_line_spacing_rule in ('单倍行距','双倍行距','1.5倍行距'):
                #     para_setting_list.append(self.style_para_line_spacing_rule)
                # else:
                #     para_setting_list.append(self.style_para_line_spacing+'倍行距')

            # if self.style_para_firstchardropcap !='' and self.style_para_firstchardropcaplines !='': 
                # para_setting_list.append('首字'+self.style_para_firstchardropcap+self.style_para_firstchardropcaplines+'磅')
            if self.style_page_break_before==True: para_setting_list.append('段前分页')
            if self.style_keep_with_next==True: para_setting_list.append('与下段同页')
            if self.style_keep_together==True: para_setting_list.append('段中不分页')
            if self.style_widow_control==True: para_setting_list.append('孤行控制')

            if len(para_setting_list)>0:
                style_des_add.append('其段落格式设置成'+'、'.join(para_setting_list))
            desc_all = '为段落【'+self.para_text_simple(self.style_text)+'】'\
                +style_des+'('+'，'.join(style_des_add)+')'+'.'
            return desc_all
        else:
            return ''
    style_description.short_description = '样式设置描述'

    def image_description(self):
        if self.image_op:
            image_desc = '在段落【'+self.para_text_simple(self.image_text)+'】后以"'+\
                self.image_position_style+'"格式插入图片"'+self.upload_image_file.name.split('/')[-1]+'"'
            if self.image_width !='':
                image_desc = image_desc + '、宽度设置为'+self.image_width+'厘米'
            if self.image_height !='': 
                image_desc = image_desc + '、高度设置为'+self.image_height+'厘米'
            return image_desc+'.'
        else:
            return ''
    image_description.short_description = '图片插入描述'

    def table_description(self):
        if self.table_op:
            table_desc = '将内容【'+self.para_text_simple(self.table_text)+'】生成表格'
            if self.table_alignment!='':
                table_desc = table_desc +'， 表格'+ [x[1] for x in TABLE_ALIGNMENT_CHOICES if x[0]==self.table_alignment][0]
            if self.table_style!='':
                table_desc = table_desc +'， 应用"'+ [x[1] for x in TABLE_STYLE_CHOICES if x[0]==self.table_style][0] +'"表格样式'
            if self.table_autofit:
                table_desc = table_desc +'，根据内容自动调整表格' 
            return table_desc+'.'
        else:
            return ''
    table_description.short_description = '表格插入描述'


    def clean(self):
        if not (self.char_edit_op or self.font_op or self.paraformat_op or self.style_op or self.image_op or self.table_op):
            raise ValidationError(_('至少选择一个操作考查'))

        # if self.table_op and (self.char_edit_op or self.font_op or self.paraformat_op or self.style_op or self.image_op):
        #     raise ValidationError(_('插入表格不能和其他操作同时选择'))

        # if (self.font_op or self.paraformat_op) and self.style_op:
        #     raise ValidationError(_('样式和单独的字体和段落格式不应同时设置'))

        # check word file
        # print('file:', self.word_question.upload_docx.upload.path)
        # try:
        #     document = Document(self.upload_docx.upload.path)
        # except:
        #     raise ValidationError({'word_question':_('文件异常')})

        # all_paras  = document.paragraphs
        # para_text = self.para_text.strip()
        # matched = 0
        # for para in para_text.split('\n'):
        #     for i in range(len(all_paras)):
        #         if all_paras[i].text.strip()==para.strip():
        #             matched = matched + 1
        # if matched != len(para_text.split('\n')):
        #     raise ValidationError({'para_text':_('未找到文章中对应内容，请检查是否输入正确内容')})

        error_dict = {}

        if self.char_edit_op and (self.char_edit_text==''):
            error_dict['char_edit_text'] = _('不能为空')
        if self.char_edit_op and (self.char_edit_origin==''):
            error_dict['char_edit_origin'] = _('不能为空')
        if self.char_edit_op and (self.char_edit_replace==''):
            error_dict['char_edit_replace'] = _('不能为空')


        if self.font_op and (self.font_text==''):
            error_dict['font_text'] = _('不能为空')

        if self.font_op and \
        (self.font_name_ascii=='' and 
        self.font_name_chinese=='' and 
        self.font_size=='' and 
        self.font_underline=='' and 
        self.font_color=='' and
        self.font_bold==False and
        self.font_italic==False):
            error_dict['font_name_chinese'] = _('至少设定一个字体相关设置')
            error_dict['font_name_ascii'] = _('')
            error_dict['font_size'] = _('')
            error_dict['font_underline'] = _('')
            error_dict['font_color'] = _('')


        if self.paraformat_op and (self.paraformat_text==''):
            error_dict['paraformat_text'] = _('不能为空')

        if self.paraformat_op and \
        (self.para_alignment=='' and 
        self.para_left_indent=='' and 
        self.para_right_indent=='' and 
        self.para_first_line_indent=='' and 
        self.para_first_line_indent_size=='' and 
        self.para_space_before=='' and 
        self.para_space_after=='' and 
        self.para_line_spacing_rule=='' and 
        self.para_line_spacing=='' and 
        # self.para_firstchardropcap=='' and 
        # self.para_firstchardropcaplines=='' and 
        self.page_break_before==False and
        self.keep_with_next==False and
        self.keep_together==False and
        self.widow_control==False):
            error_dict['para_alignment'] = _('至少选择一个段落格式相关设置')
            error_dict['para_left_indent'] = _('')
            error_dict['para_right_indent'] = _('')
            error_dict['para_first_line_indent'] = _('')
            error_dict['para_first_line_indent_size'] = _('')
            error_dict['para_space_before'] = _('')
            error_dict['para_space_after'] = _('')
            error_dict['para_line_spacing_rule'] = _('')
            error_dict['para_line_spacing'] = _('')
            # error_dict['para_firstchardropcap'] = _('')
            # error_dict['para_firstchardropcaplines'] = _('')

        if self.paraformat_op and (self.para_first_line_indent=='' and self.para_first_line_indent_size!=''):
            error_dict['para_first_line_indent'] = _('不能为空')
            error_dict['para_first_line_indent_size'] = _('*')
        if self.paraformat_op and (self.para_first_line_indent !='' and self.para_first_line_indent_size==''):
            error_dict['para_first_line_indent'] = _('*')
            error_dict['para_first_line_indent_size'] = _('不能为空')

        if self.paraformat_op and (self.para_line_spacing_rule=='' and self.para_line_spacing!=''):
            error_dict['para_line_spacing_rule'] = _('不能为空')
            error_dict['para_line_spacing'] = _('*')
        if self.paraformat_op and (self.para_line_spacing_rule =='MULTIPLE (5)' and self.para_line_spacing==''):
            error_dict['para_line_spacing_rule'] = _('*')
            error_dict['para_line_spacing'] = _('不能为空')

        # if self.paraformat_op and (self.para_firstchardropcap=='' and self.para_firstchardropcaplines!=''):
        #     error_dict['para_firstchardropcap'] = _('不能为空')
        #     error_dict['para_firstchardropcaplines'] = _('*')
        # if self.paraformat_op and (self.para_firstchardropcap !='' and self.para_firstchardropcaplines==''):
        #     error_dict['para_firstchardropcap'] = _('*')
        #     error_dict['para_firstchardropcaplines'] = _('不能为空')


        if self.style_op and self.style_name=='':
            error_dict['style_name'] = _('内容不能为空')

        if self.style_op and self.style_text=='':
            error_dict['style_text'] = _('内容不能为空')

        if self.style_op and self.style_name in ['新样式1', '新样式2'] and \
               (self.style_font_name_ascii=='' and 
                self.style_font_name_chinese=='' and 
                self.style_font_size=='' and 
                self.style_font_underline=='' and 
                self.style_font_color=='' and
                self.style_font_bold==False and
                self.style_font_italic==False and
                self.style_para_alignment=='' and 
                self.style_para_left_indent=='' and 
                self.style_para_right_indent=='' and 
                self.style_para_first_line_indent=='' and 
                self.style_para_first_line_indent_size=='' and 
                self.style_para_space_before=='' and 
                self.style_para_space_after=='' and 
                self.style_para_line_spacing_rule=='' and 
                self.style_para_line_spacing=='' and 
                self.style_page_break_before==False and
                self.style_keep_with_next==False and
                self.style_keep_together==False and
                self.style_widow_control==False
                ):
            error_dict['style_name'] = _('至少为新样式设定一个具体设置')
            error_dict['style_font_name_chinese'] = _('')
            error_dict['style_font_name_ascii'] = _('')
            error_dict['style_font_size'] = _('')
            error_dict['style_font_underline'] = _('')
            error_dict['style_font_color'] = _('')
            error_dict['style_para_alignment'] = _('')
            error_dict['style_para_left_indent'] = _('')
            error_dict['style_para_right_indent'] = _('')
            error_dict['style_para_first_line_indent'] = _('')
            error_dict['style_para_first_line_indent_size'] = _('')
            error_dict['style_para_space_before'] = _('')
            error_dict['style_para_space_after'] = _('')
            error_dict['style_para_line_spacing_rule'] = _('')
            error_dict['style_para_line_spacing'] = _('')

        if self.style_op and \
        (self.style_para_first_line_indent=='' and 
         self.style_para_first_line_indent_size!=''):
            error_dict['style_para_first_line_indent'] = _('不能为空')
            error_dict['style_para_first_line_indent_size'] = _('*')

        if self.style_op and \
            (self.style_para_first_line_indent !='' and 
             self.style_para_first_line_indent_size==''):
            error_dict['style_para_first_line_indent'] = _('*')
            error_dict['style_para_first_line_indent_size'] = _('不能为空')

        if self.style_op and (self.style_para_line_spacing_rule=='' and self.style_para_line_spacing!=''):
            error_dict['style_para_line_spacing'] = _('*')
            error_dict['style_para_line_spacing_rule'] = _('不能为空')
        if self.style_op and (self.style_para_line_spacing_rule !='' and self.style_para_line_spacing==''):
            error_dict['style_para_line_spacing'] = _('不能为空')
            error_dict['style_para_line_spacing_rule'] = _('*')

        # print('self.upload_image_file',  self.upload_image_file.name)
        if self.image_op:
            if self.upload_image_file.name=='':
                error_dict['upload_image_file'] = _('必须上传图片(jpg,bmp,png)')
            if self.image_text=='':
                error_dict['image_text'] = _('不能为空')
        
        if self.table_op:
            if self.table_text=='':
                error_dict['table_text'] = _('不能为空')

        if len(error_dict)>0:
            raise ValidationError(error_dict)

    def find_text_para_index(self, all_paras_, para_text_list_):
        matched_index_list = []
        i = 0
        for para in para_text_list_:
            while i < (len(all_paras_)):
                if Levenshtein.ratio(all_paras_[i].text, para)>0.6:
                    matched_index_list.append(i)
                    break
                i = i + 1
        return matched_index_list

    def compare_operation(self, docx_test):
        all_paras  = docx_test.paragraphs
        result_list = []

        # 
        if self.char_edit_op:
            para_text_list = self.char_edit_text.strip().split('\n')
            matched_list = self.find_text_para_index(all_paras, para_text_list)
            if len(matched_list) != len(para_text_list):
                result_list.append(['文字编辑', '未找到文章中对应段落', '']) 
            else:
                count_origin, count_replace = 0, 0
                count_origin = self.char_edit_text.count(self.char_edit_origin)
                for ind in  matched_list:
                    count_replace += all_paras[ind].text.count(self.char_edit_replace)
                result_list.append(['替换',str(count_origin),str(count_replace)])

        if self.font_op:
            para_text_list = self.font_text.strip().split('\n')
            matched_list = self.find_text_para_index(all_paras, para_text_list)
            if len(matched_list) != len(para_text_list):
                result_list.append(['字体设置', '未找到文章中对应段落', '']) 
            else:
                r0 = all_paras[matched_list[0]].runs[0]
                pstyle_font = all_paras[matched_list[0]].style.font
                if self.font_name_chinese !='': 
                    result_list.append(['中文', self.font_name_chinese, r0.font.name_eastasia or pstyle_font.name_eastasia, r0.font.name_eastasia, pstyle_font.name_eastasia])
                if self.font_name_ascii !='':  
                    result_list.append(['西文', self.font_name_ascii, r0.font.name or pstyle_font.name, r0.font.name, pstyle_font.name])
                if self.font_size !='': 
                    result_list.append(['字号', self.font_size, getPtorNone(r0.font.size) or getPtorNone(pstyle_font.size) ,getPtorNone(r0.font.size), getPtorNone(pstyle_font.size)])
                if self.font_color !='': 
                    result_list.append(['字色', self.font_color, str(r0.font.color.rgb or pstyle_font.color.rgb), str(r0.font.color.rgb), str(pstyle_font.color.rgb)])
                if self.font_bold==True: 
                    result_list.append(['粗体', self.font_bold, r0.font.bold or pstyle_font.bold, r0.font.bold, pstyle_font.bold])
                if self.font_italic==True: 
                    result_list.append(['斜体', self.font_italic, r0.font.italic or pstyle_font.italic, r0.font.italic, pstyle_font.italic])
                if self.font_underline !='': 
                    result_list.append(['下划线', self.font_underline, str(r0.font.underline or pstyle_font.underline) ,str(r0.font.underline), str(pstyle_font.underline)])

        if self.paraformat_op:
            para_text_list = self.paraformat_text.strip().split('\n')
            matched_list = self.find_text_para_index(all_paras, para_text_list)
            if len(matched_list) != len(para_text_list):
                result_list.append(['段落设置', '未找到文章中对应段落', '']) 
            else:
                p_format = all_paras[matched_list[0]].paragraph_format
                pstyle_format = all_paras[matched_list[0]].style.paragraph_format

                if self.para_alignment !='': 
                    para_alignment = PARA_ALIGNMENT_CHOICES[0][0]  ## 'LEFT (0)'
                    if p_format.alignment is not None:
                        para_alignment = p_format.alignment
                    elif pstyle_format.alignment is not None:
                        para_alignment = pstyle_format.alignment
                    result_list.append(['对齐', self.para_alignment, str(para_alignment), str(pstyle_format.alignment)])

                if self.para_left_indent !='': 
                    result_list.append(['左缩进', 
                    (self.para_left_indent), getSzorNone(p_format.left_indent, pstyle_format.left_indent), 
                    getPtorNone(p_format.left_indent), 
                    getPtorNone(pstyle_format.left_indent)])
                if self.para_right_indent !='': 
                    result_list.append(['右缩进', 
                    (self.para_left_indent), getSzorNone(p_format.right_indent, pstyle_format.right_indent), 
                    getPtorNone(p_format.right_indent), 
                    getPtorNone(pstyle_format.right_indent)])

                if self.para_first_line_indent !='' and self.para_first_line_indent_size !='':
                    if self.para_first_line_indent == PARA_FIRST_LINE_INDENT_CHOICES[0][0]:  ## +
                        result_list.append(['首行缩进', (self.para_first_line_indent_size), 
                        getSzorNone(p_format.first_line_indent, pstyle_format.first_line_indent),
                            p_format.first_line_indent, pstyle_format.first_line_indent])
                    else:  ## -
                        result_list.append(['首行缩进', (self.para_first_line_indent_size), 
                        -getSzorNone(p_format.first_line_indent, pstyle_format.first_line_indent),
                            p_format.first_line_indent, pstyle_format.first_line_indent])

                if self.para_space_before !='': 
                    result_list.append(['段前', (self.para_space_before), 
                    getSzorNone(p_format.space_before, pstyle_format.space_before), 
                    getPtorNone(p_format.space_before), 
                    getPtorNone(pstyle_format.space_before)])
                if self.para_space_after !='':  
                    result_list.append(['段后', (self.para_space_after), 
                    getSzorNone(p_format.space_after, pstyle_format.space_after), 
                    getPtorNone(p_format.space_after), 
                    getPtorNone(pstyle_format.space_after)])

                # if self.para_firstchardropcap !='' and self.para_firstchardropcaplines !='': 
                #     result_list.append(['首字下沉', self.para_firstchardropcap, str(p_format.first_char_dropcap)])
                #     result_list.append(['首字下沉行数', 
                #         self.para_firstchardropcaplines, str(p_format.first_char_dropcap_lines)])

                if self.para_line_spacing_rule !='': 
                    result_list.append(['行间距规则', self.para_line_spacing_rule, str(p_format.line_spacing_rule or pstyle_format.line_spacing_rule), \
                        str(pstyle_format.line_spacing_rule)])
                    if self.para_line_spacing_rule not in (x[0] for x in PARA_LINE_SPACING_RULE_CHOICES[:3]):
                        result_list.append(['行距', (self.para_line_spacing), getLineorNone(p_format.line_spacing, pstyle_format.line_spacing), 
                        pstyle_format.line_spacing])

                if self.page_break_before==True: 
                    result_list.append(['段前分页', self.page_break_before, p_format.page_break_before or pstyle_format.page_break_before])
                if self.keep_with_next==True: 
                    result_list.append(['与下段同页',self.keep_with_next, p_format.keep_with_next or pstyle_format.keep_with_next])
                if self.keep_together==True: 
                    result_list.append(['段中不分页',self.keep_together, p_format.keep_together or pstyle_format.keep_together])
                if self.widow_control==True:
                    widow_control = True
                    if p_format.widow_control==False or (p_format.widow_control is None and pstyle_format.widow_control==False):
                        widow_control = False
                    result_list.append(['孤行控制', True, widow_control, p_format.widow_control , pstyle_format.widow_control])

        if self.image_op:
            all_images = docx_test.inline_shapes
            if len(all_images)>0:
                result_list.append(['插入图片','True','True',])
                if self.image_width !='':
                    result_list.append(['图片宽度', (self.image_width), round(all_images[0].width.cm),])
                if self.image_height !='': 
                    result_list.append(['图片高度', (self.image_height), round(all_images[0].height.cm),])
            else:
                result_list.append(['插入图片','True','False',])
                if self.image_width !='':
                    result_list.append(['图片宽度', (self.image_width), 0,])
                if self.image_height !='': 
                    result_list.append(['图片高度', (self.image_height), 0,])

        if  self.table_op:
            all_tables = docx_test.tables
            if len(all_tables):
                result_list.append(['插入表格','True','True',])
                if self.table_autofit:
                    result_list.append(['表格自动调整', 'True', str(all_tables[0].autofit),])
                if self.table_alignment!='':
                    result_list.append(['表格对齐', self.table_alignment, str(all_tables[0].alignment),])
                if self.table_style!='':
                    result_list.append(['表格样式', self.table_style, str(all_tables[0].style.name),])
            else:
                result_list.append(['插入表格','True','False',])
                if self.table_autofit:
                    result_list.append(['表格自动调整', 'True', '--',])
                if self.table_alignment!='':
                    result_list.append(['表格对齐', self.table_alignment, '--',])
                if self.table_style!='':
                    result_list.append(['表格样式', self.table_style, '--',])


        if self.style_op:
            para_text_list = self.style_text.strip().split('\n')
            matched_list = self.find_text_para_index(all_paras, para_text_list)
            if len(matched_list) != len(para_text_list):
                result_list.append(['样式设置', '未找到文章中对应段落', '']) 
            else:
                r0 = all_paras[matched_list[0]].runs[0]
                pstyle_font = all_paras[matched_list[0]].style.font
                p_format = all_paras[matched_list[0]].paragraph_format
                pstyle_format = all_paras[matched_list[0]].style.paragraph_format

                result_list.append(['样式名称', self.style_name, all_paras[matched_list[0]].style.name])

                if self.style_font_name_chinese !='': 
                    result_list.append(['中文', self.style_font_name_chinese, pstyle_font.name_eastasia])
                if self.style_font_name_ascii !='':  
                    result_list.append(['西文', self.style_font_name_ascii, pstyle_font.name, r0.font.name, pstyle_font.name])
                if self.style_font_size !='': 
                    result_list.append(['字号', self.style_font_size, getPtorNone(pstyle_font.size) ,getPtorNone(r0.font.size), getPtorNone(pstyle_font.size)])
                if self.style_font_color !='': 
                    result_list.append(['字色', self.style_font_color, str(pstyle_font.color.rgb), str(r0.font.color.rgb)])
                if self.style_font_bold==True: 
                    result_list.append(['粗体', self.style_font_bold, pstyle_font.bold, r0.font.bold, pstyle_font.bold])
                if self.style_font_italic==True: 
                    result_list.append(['斜体', self.style_font_italic, pstyle_font.italic, r0.font.italic, pstyle_font.italic])
                if self.style_font_underline !='': 
                    result_list.append(['下划线', self.style_font_underline, str(pstyle_font.underline) ,str(r0.font.underline), str(pstyle_font.underline)])

                if self.style_para_alignment !='': 
                    # para_alignment = PARA_ALIGNMENT_CHOICES[0][0]  ## 'LEFT (0)'
                    # if pstyle_format.alignment:
                        # para_alignment = pstyle_format.alignment
                    result_list.append(['对齐', self.style_para_alignment, str(pstyle_format.alignment or PARA_ALIGNMENT_CHOICES[0][0])])

                if self.style_para_left_indent !='': 
                    result_list.append(['左缩进', 
                    (self.style_para_left_indent),  
                    getPtorNone(pstyle_format.left_indent)])
                if self.style_para_right_indent !='': 
                    result_list.append(['右缩进', 
                    (self.style_para_left_indent), 
                    getPtorNone(pstyle_format.right_indent)])

                if self.style_para_first_line_indent !='' and self.style_para_first_line_indent_size !='':
                    if self.style_para_first_line_indent == PARA_FIRST_LINE_INDENT_CHOICES[0][0]:  ## +
                        result_list.append(['首行缩进', (self.style_para_first_line_indent_size), 
                            getPtorNone(pstyle_format.first_line_indent)])
                    else:  ## -
                        result_list.append(['首行缩进', (self.style_para_first_line_indent_size), 
                            -getPtorNone(pstyle_format.first_line_indent)])

                if self.style_para_space_before !='': 
                    result_list.append(['段前', (self.style_para_space_before), 
                    getPtorNone(pstyle_format.space_before)])
                if self.style_para_space_after !='':  
                    result_list.append(['段后', (self.style_para_space_after), 
                    getPtorNone(pstyle_format.space_after)])

                # if self.para_firstchardropcap !='' and self.para_firstchardropcaplines !='': 
                #     result_list.append(['首字下沉', self.para_firstchardropcap, str(p_format.first_char_dropcap)])
                #     result_list.append(['首字下沉行数', 
                #         self.para_firstchardropcaplines, str(p_format.first_char_dropcap_lines)])

                if self.style_para_line_spacing_rule !='': 
                    result_list.append(['行间距规则', self.style_para_line_spacing_rule, str(p_format.line_spacing_rule or pstyle_format.line_spacing_rule), \
                        str(pstyle_format.line_spacing_rule)])
                    if self.style_para_line_spacing_rule not in (x[0] for x in PARA_LINE_SPACING_RULE_CHOICES[:3]):
                        result_list.append(['行距', (self.style_para_line_spacing), getLineorNone(None, pstyle_format.line_spacing), 
                        pstyle_format.line_spacing])

                if self.style_page_break_before==True: 
                    result_list.append(['段前分页', self.style_style_page_break_before, pstyle_format.page_break_before])
                if self.style_keep_with_next==True: 
                    result_list.append(['与下段同页',self.style_keep_with_next, pstyle_format.keep_with_next])
                if self.style_keep_together==True: 
                    result_list.append(['段中不分页',self.style_keep_together, pstyle_format.keep_together])
                if self.style_widow_control==True:
                    result_list.append(['孤行控制', True, pstyle_format.widow_control])

        return result_list
        
    def test_(self):
        if self.upload_docx is not None:
            # check valid docx file
            try:
                document_test = Document(self.upload_docx.path)
            except:
                return self.upload_docx.path+'打开异常'

            result_list = self.compare_operation(document_test)
            print(result_list)

            return format_html("<ol>") + \
                format_html_join(
                '\n', '<li style="color:black;">{} {} {}</li>',
                (x for x in result_list)
                ) \
                + format_html("</ol>")
        else:
            return '没有上传测试文件'
    test_.short_description = '测试结果'

    def score_(self, file_docx):
        if file_docx:
            # check valid docx file
            try:
                document_test = Document(file_docx)
            except:
                return 0

            result_list = self.compare_operation(document_test)
            # print(result_list)
            scores = [(len(x)>2 and x[1]==x[2])*1 for x in result_list]
            # print(scores)
            s = 0
            for i in scores:
                s = s + i

            return s*1.0/len(scores)

    def docx_path_(self):
        if self.upload_docx:
            return self.upload_docx.path
        else:
            return 'None'
    docx_path_.short_description = 'Word文件地址'


    #########
    pub_date = models.DateTimeField('创建时间')

    upload_docx = models.FileField(verbose_name='上传word文件(.docx)',upload_to='upload_docx/',
        validators=[validate_docx])

    ############## 文字查找替换
    char_edit_text = models.TextField('文字编辑要考查的文字内容', blank=True, default='')
    char_edit_op = models.BooleanField('考查文字查找替换？', default=False)
    char_edit_origin = models.CharField('原词', max_length=200, blank=True, default='')
    char_edit_replace = models.CharField('替换为', max_length=200, blank=True, default='')

    ############## 字体设置
    font_text = models.TextField('字体要考查的文字内容', blank=True, default='')
    font_op = models.BooleanField('考查字体设置？', default=False)
    font_name_ascii = models.CharField('英文字体', choices=FONT_NAME_ASCII_CHOICES, max_length=200, blank=True, default='') 
    font_name_chinese = models.CharField('中文字体', choices=FONT_NAME_CHINESE_CHOICES, max_length=200, blank=True, default='')
    font_size = models.CharField('字号(磅)', choices=FONT_SIZE_CHOICES, max_length=200, blank=True, default='')
    font_bold = models.BooleanField('粗体', blank=True, default='')
    font_italic = models.BooleanField('斜体', blank=True, default='')
    font_underline = models.CharField('下划线', choices=FONT_UNDERLINE_CHOICES, max_length=200, blank=True, default='')
    font_color = models.CharField('字色(标准色)', choices=FONT_COLOR_CHOICES, max_length=200, blank=True, default='')

    ############## 段落格式设置
    paraformat_text = models.TextField('段落格式要考查的文字内容', blank=True, default='')
    paraformat_op = models.BooleanField('考查段落格式设置？', default=False)
    para_alignment = models.CharField('段落对齐', choices=PARA_ALIGNMENT_CHOICES, max_length=200, blank=True, default='')
    para_left_indent = models.CharField('左侧缩进(磅)', choices=INDENT_NUM_CHOICES, max_length=200, blank=True, default='')
    para_right_indent = models.CharField('右侧缩进(磅)', choices=INDENT_NUM_CHOICES, max_length=200, blank=True, default='')
    para_first_line_indent = models.CharField('首行缩进', choices=PARA_FIRST_LINE_INDENT_CHOICES, max_length=200, blank=True, default='')
    para_first_line_indent_size = models.CharField('首行缩进距离(磅)', choices=INDENT_NUM_CHOICES, max_length=200, blank=True, default='')

    para_space_before = models.CharField('段前间距(磅)',  choices=INDENT_NUM_CHOICES, max_length=200, blank=True, default='')
    para_space_after = models.CharField('段后间距(磅)',  choices=INDENT_NUM_CHOICES, max_length=200, blank=True, default='')
    para_line_spacing_rule = models.CharField('行距规则', choices=PARA_LINE_SPACING_RULE_CHOICES, max_length=200, blank=True, default='')
    para_line_spacing = models.CharField('行距(行)',  choices=LINE_NUM_CHOICES, max_length=200, blank=True, default='')

    # para_firstchardropcap = models.CharField('首字下沉', choices=PARA_FIRSTCHARDROPCAP_CHOICES, max_length=200, blank=True, default='')
    # para_firstchardropcaplines = models.CharField('下沉(行)',  choices=LINE_NUM_CHOICES, max_length=200, blank=True, default='')

    page_break_before = models.BooleanField('段前分页', default=False)
    keep_with_next = models.BooleanField('与下段同页', default=False)
    keep_together = models.BooleanField('段中不分页', default=False)
    widow_control = models.BooleanField('孤行控制', default=False)

    ############## 样式设置
    style_text = models.TextField('样式要考查的文字内容', blank=True, default='')
    style_op = models.BooleanField('考查样式设置？', default=False)
    style_name = models.CharField('样式名称', choices=STYLE_NAME_CHOICES, max_length=200, blank=True, default='') 
    style_font_name_ascii = models.CharField('英文字体', choices=FONT_NAME_ASCII_CHOICES, max_length=200, blank=True, default='') 
    style_font_name_chinese = models.CharField('中文字体', choices=FONT_NAME_CHINESE_CHOICES, max_length=200, blank=True, default='')
    style_font_size = models.CharField('字号', choices=FONT_SIZE_CHOICES, max_length=200, blank=True, default='')
    style_font_bold = models.BooleanField('粗体', default=False)
    style_font_italic = models.BooleanField('斜体', default=False)
    style_font_underline = models.CharField('下划线', choices=FONT_UNDERLINE_CHOICES, max_length=200, blank=True, default='')
    style_font_color = models.CharField('字色(标准色)', choices=FONT_COLOR_CHOICES, max_length=200, blank=True, default='')

    style_para_alignment = models.CharField('段落对齐', choices=PARA_ALIGNMENT_CHOICES, max_length=200, blank=True, default='')
    style_para_left_indent = models.CharField('左侧缩进(磅)',  choices=INDENT_NUM_CHOICES, max_length=200, blank=True, default='')
    style_para_right_indent = models.CharField('右侧缩进(磅)',  choices=INDENT_NUM_CHOICES, max_length=200, blank=True, default='')
    style_para_first_line_indent = models.CharField('首行缩进', choices=PARA_FIRST_LINE_INDENT_CHOICES, max_length=200, blank=True, default='')
    style_para_first_line_indent_size = models.CharField('首行缩进距离(磅)',  choices=INDENT_NUM_CHOICES, max_length=200, blank=True, default='')

    style_para_space_before = models.CharField('段前间距(磅)',   choices=INDENT_NUM_CHOICES, max_length=200, blank=True, default='')
    style_para_space_after = models.CharField('段后间距(磅)',   choices=INDENT_NUM_CHOICES, max_length=200, blank=True, default='')
    style_para_line_spacing_rule = models.CharField('行距规则', choices=PARA_LINE_SPACING_RULE_CHOICES, max_length=200, blank=True, default='')
    style_para_line_spacing = models.CharField('行距(行)',   choices=LINE_NUM_CHOICES, max_length=200, blank=True, default='')

    style_page_break_before = models.BooleanField('段前分页', default=False)
    style_keep_with_next = models.BooleanField('与下段同页', default=False)
    style_keep_together = models.BooleanField('段中不分页', default=False)
    style_widow_control = models.BooleanField('孤行控制', default=False)

    ############## 图片插入
    image_text = models.TextField('图片插入的文字内容', blank=True, default='')
    image_op = models.BooleanField('考查图片插入？', default=False)
    image_position_style = models.CharField('位置类型', choices=IMAGE_POSITION_STYLE_CHOICES, max_length=20, default='嵌入文本行中')
    image_width = models.CharField('图像宽(厘米)',   choices=INDENT_NUM_CHOICES, max_length=200, blank=True, default='')
    image_height = models.CharField('图像高(厘米)',   choices=INDENT_NUM_CHOICES, max_length=200, blank=True, default='')
    upload_image_file = models.FileField(upload_to='upload_image/', null=True, blank=True, 
    validators=[validate_image], verbose_name='上传图片')

    ############## 表格插入
    table_text = models.TextField('表格要考查的文字内容', blank=True, default='')
    table_op = models.BooleanField('考查表格插入？', default=False)
    table_alignment = models.CharField('对齐方式', choices=TABLE_ALIGNMENT_CHOICES, max_length=10, default='')
    table_autofit = models.BooleanField('表格宽度自动调整？', default=False)
    table_style = models.CharField('表格样式', choices=TABLE_STYLE_CHOICES, max_length=30, default='')


# # Create your models here.
# class WordOperations(models.Model):

#     class Meta:
#         verbose_name = 'Word操作列表'
#         verbose_name_plural = 'Word操作列表'

#     def __str__(self):
#         return 'Word操作'+str(self.id)

#     def word_question_info(self):
#         if self.word_question is not None:
#             return self.word_question.__str__()
#     word_question_info.short_description = '操作所属题目'

#     def para_text_simple(self):
#         if len(self.para_text)<30:
#             return self.para_text.strip()
#         else:
#             return self.para_text[:15] +' ...... '+self.para_text[-15:]
#     para_text_simple.short_description = '考查段落内容'

#     def operations_list(self):
#             op_list = []
#             if self.char_edit_op:
#                 op_list.append('查找替换')
#             if self.font_op:
#                 op_list.append('字体设置')
#             if self.paraformat_op:
#                 op_list.append('段落设置')
#             if self.style_op:
#                 op_list.append('样式设置')
#             if self.image_op:
#                 op_list.append('图片插入')
#             return '/'.join(op_list)
#     operations_list.short_description = '涉及操作'

#     def compare_operation(self, document_test):
#         all_paras  = document_test.paragraphs
#         para_text_list = self.para_text.strip().split('\n')
#         matched_list = []
#         for para in para_text_list:
#             for i in range(len(all_paras)):
#                 if all_paras[i].text.strip()==para.strip():
#                     matched_list.append(i)
#         if len(matched_list) != len(para_text_list):
#             return ['red', '未找到文章中对应考查段落']

#         result_list = []
#         # 
#         if self.char_edit_op:
#             count_origin, count_replace = 0, 0
#             for para in para_text_list:
#                 count_origin += para.count(self.char_edit_origin)
#             for ind in  matched_list:
#                 count_replace += all_paras[ind].text.count(self.char_edit_replace)
#             result_list.append(['替换',str(count_origin),str(count_replace)])

#         if self.font_op:
#             r0 = all_paras[matched_list[0]].runs[0]
#             pstyle_font = all_paras[matched_list[0]].style.font
#             if self.font_name_chinese !='': 
#                 result_list.append(['中文', self.font_name_chinese, r0.font.name_eastasia or pstyle_font.name_eastasia, r0.font.name_eastasia, pstyle_font.name_eastasia])
#             if self.font_name_ascii !='':  
#                 result_list.append(['西文', self.font_name_ascii, r0.font.name or pstyle_font.name, r0.font.name, pstyle_font.name])
#             if self.font_size !='': 
#                 result_list.append(['字号', (self.font_size), getPtorNone(r0.font.size) or getPtorNone(pstyle_font.size) ,getPtorNone(r0.font.size), getPtorNone(pstyle_font.size)])
#             if self.font_color !='': 
#                 result_list.append(['字色', self.font_color, str(r0.font.color.rgb or pstyle_font.color.rgb), str(r0.font.color.rgb), str(pstyle_font.color.rgb)])
#             if self.font_bold==True: 
#                 result_list.append(['粗体', self.font_bold, r0.font.bold or pstyle_font.bold, r0.font.bold, pstyle_font.bold])
#             if self.font_italic==True: 
#                 result_list.append(['斜体', self.font_italic, r0.font.italic or pstyle_font.italic, r0.font.italic, pstyle_font.italic])
#             if self.font_underline !='': 
#                 result_list.append(['下划线', self.font_underline, str(r0.font.underline or pstyle_font.underline) ,str(r0.font.underline), str(pstyle_font.underline)])

#         if self.paraformat_op:
#             p_format = all_paras[matched_list[0]].paragraph_format
#             pstyle_format = all_paras[matched_list[0]].style.paragraph_format

#             if self.para_alignment !='': 
#                 para_alignment = PARA_ALIGNMENT_CHOICES[0][0]  ## 'LEFT (0)'
#                 if p_format.alignment is not None:
#                     para_alignment = p_format.alignment
#                 elif pstyle_format.alignment is not None:
#                     para_alignment = pstyle_format.alignment

#                 result_list.append(['对齐', self.para_alignment, str(para_alignment), str(pstyle_format.alignment)])
#             if self.para_left_indent !='': 
#                 result_list.append(['左缩进', 
#                 (self.para_left_indent), getSzorNone(p_format.left_indent, pstyle_format.left_indent), 
#                 getPtorNone(p_format.left_indent), 
#                 getPtorNone(pstyle_format.left_indent)])
#             if self.para_right_indent !='': 
#                 result_list.append(['右缩进', 
#                 (self.para_left_indent), getSzorNone(p_format.right_indent, pstyle_format.right_indent), 
#                 getPtorNone(p_format.right_indent), 
#                 getPtorNone(pstyle_format.right_indent)])

#             if self.para_first_line_indent !='' and self.para_first_line_indent_size !='':
#                 if self.para_first_line_indent == PARA_FIRST_LINE_INDENT_CHOICES[0][0]:  ## +
#                     result_list.append(['首行缩进', (self.para_first_line_indent_size), 
#                     getSzorNone(p_format.first_line_indent, pstyle_format.first_line_indent),
#                         p_format.first_line_indent, pstyle_format.first_line_indent])
#                 else:  ## -
#                     result_list.append(['首行缩进', (self.para_first_line_indent_size), 
#                     -getSzorNone(p_format.first_line_indent, pstyle_format.first_line_indent),
#                         p_format.first_line_indent, pstyle_format.first_line_indent])

#             if self.para_space_before !='': 
#                 result_list.append(['段前', (self.para_space_before), 
#                 getSzorNone(p_format.space_before, pstyle_format.space_before), 
#                 getPtorNone(p_format.space_before), 
#                 getPtorNone(pstyle_format.space_before)])
#             if self.para_space_after !='':  
#                 result_list.append(['段后', (self.para_space_after), 
#                 getSzorNone(p_format.space_after, pstyle_format.space_after), 
#                 getPtorNone(p_format.space_after), 
#                 getPtorNone(pstyle_format.space_after)])

#             # if self.para_firstchardropcap !='' and self.para_firstchardropcaplines !='': 
#             #     result_list.append(['首字下沉', self.para_firstchardropcap, str(p_format.first_char_dropcap)])
#             #     result_list.append(['首字下沉行数', 
#             #         self.para_firstchardropcaplines, str(p_format.first_char_dropcap_lines)])

#             if self.para_line_spacing_rule !='': 
#                 result_list.append(['行间距规则', self.para_line_spacing_rule, str(p_format.line_spacing_rule or pstyle_format.line_spacing_rule), \
#                     str(pstyle_format.line_spacing_rule)])
#                 if self.para_line_spacing_rule not in (x[0] for x in PARA_LINE_SPACING_RULE_CHOICES[:3]):
#                     result_list.append(['行距', (self.para_line_spacing), getLineorNone(p_format.line_spacing, pstyle_format.line_spacing), 
#                     pstyle_format.line_spacing])

#             if self.page_break_before==True: 
#                 result_list.append(['段前分页', self.page_break_before, p_format.page_break_before or pstyle_format.page_break_before])
#             if self.keep_with_next==True: 
#                 result_list.append(['与下段同页',self.keep_with_next, p_format.keep_with_next or pstyle_format.keep_with_next])
#             if self.keep_together==True: 
#                 result_list.append(['段中不分页',self.keep_together, p_format.keep_together or pstyle_format.keep_together])
#             if self.widow_control==True:
#                 widow_control = True
#                 if p_format.widow_control==False or (p_format.widow_control is None and pstyle_format.widow_control==False):
#                     widow_control = False
#                 result_list.append(['孤行控制', True, widow_control, p_format.widow_control , pstyle_format.widow_control])

#         if self.image_op:
#             all_images = document_test.inline_shapes
#             if len(all_images)>0:
#                 result_list.append(['插入图片','True','True',])
#                 if self.image_width !='':
#                     result_list.append(['图片宽度', (self.image_width), int(all_images[0].width.cm)])
#                     print(int(all_images[0].width.cm))
#                 if self.image_height !='': 
#                     result_list.append(['图片高度', (self.image_height), int(all_images[0].height.cm)])
#             else:
#                 result_list.append(['插入图片','True','False',])
#                 if self.image_width !='':
#                     result_list.append(['图片宽度', (self.image_width), 0,])
#                 if self.image_height !='': 
#                     result_list.append(['图片高度', (self.image_height), 0,])

#         return result_list

#     def operation_description_all(self):
#         description_list = [self.char_edit_description(), self.font_description(),
#         self.paraformat_description(), self.style_description()]
#         description_list = [x for x in description_list if len(x)>0]
#         op_desc_all = ''
#         if len(description_list): 
#             op_desc_all = '将段落【'+self.para_text_simple()+'】'+'，'.join(description_list)+'。'
#         if len(self.image_description()):
#             op_desc_all = op_desc_all + self.image_description() +'。'
#         if len(self.table_description()):
#             op_desc_all = op_desc_all + self.table_description() +'。'
#         return op_desc_all
#     operation_description_all.short_description = '题目文字描述'

#     def char_edit_description(self):
#         if self.char_edit_op:
#             return '所有“'+self.char_edit_origin+'”替换成“'+self.char_edit_replace+'”'
#         else:
#             return ''
#     char_edit_description.short_description = '查找替换文字描述'

#     def font_description(self):
#         if self.font_op:
#             setting_list=[]
#             if self.font_name_chinese !='': setting_list.append('中文'+self.font_name_chinese)
#             if self.font_name_ascii !='':  setting_list.append('西文'+self.font_name_ascii)
#             if self.font_size !='': setting_list.append('字号'+self.font_size+'磅')
#             if self.font_color !='': setting_list.append('标准色'+[x[1] for x in FONT_COLOR_CHOICES if x[0]==self.font_color][0])
#             if self.font_bold==True: setting_list.append('粗体')
#             if self.font_italic==True: setting_list.append('斜体')
#             if self.font_underline !='': 
#                 setting_list.append([x[1] for x in FONT_UNDERLINE_CHOICES if x[0]==self.font_underline][0])
#             return '字体设置成'+'、'.join(setting_list)
#         else:
#             return ''
#     font_description.short_description = '字体设置描述'

#     def paraformat_description(self):
#         if self.paraformat_op:
#             setting_list=[]
#             if self.para_alignment !='': 
#                 setting_list.append([x[1] for x in PARA_ALIGNMENT_CHOICES if x[0]==self.para_alignment][0])
#                 # setting_list.append(self.para_alignment)
#             if self.para_left_indent !='': setting_list.append('左缩进'+self.para_left_indent+'磅')
#             if self.para_right_indent !='': setting_list.append('右缩进'+self.para_right_indent+'磅')
#             if self.para_first_line_indent !='' and self.para_first_line_indent_size !='': 
#                 setting_list.append(self.para_first_line_indent+self.para_first_line_indent_size+'磅')
#             if self.para_space_before !='': setting_list.append('段前'+self.para_space_before+'磅')
#             if self.para_space_after !='':  setting_list.append('段后'+self.para_space_after+'磅')
#             if self.para_line_spacing_rule !='': 
#                 if self.para_line_spacing_rule in (x[0] for x in PARA_LINE_SPACING_RULE_CHOICES[:3]):
#                     setting_list.append([x[1] for x in PARA_LINE_SPACING_RULE_CHOICES if x[0]==self.para_line_spacing_rule][0])
#                 else:
#                     setting_list.append([x[1] for x in PARA_LINE_SPACING_RULE_CHOICES if x[0]==self.para_line_spacing_rule][0]+self.para_line_spacing+'倍行距')
            
#             # if self.para_firstchardropcap !='' and self.para_firstchardropcaplines !='': 
#             #     if self.para_firstchardropcap == PARA_FIRSTCHARDROPCAP_CHOICES[0][0]:
#             #         setting_list.append('首字下沉'+self.para_firstchardropcaplines+'行')
#             #     if self.para_firstchardropcap == PARA_FIRSTCHARDROPCAP_CHOICES[1][0]:
#             #         setting_list.append('首字悬挂'+self.para_firstchardropcaplines+'行')
                
#             if self.page_break_before==True: setting_list.append('段前分页')
#             if self.keep_with_next==True: setting_list.append('与下段同页')
#             if self.keep_together==True: setting_list.append('段中不分页')
#             if self.widow_control==True: setting_list.append('孤行控制')

#             return '段落格式设置成'+'、'.join(setting_list)
#         else:
#             return ''
#     paraformat_description.short_description = '段落设置描述'

#     def style_description(self):
#         if self.style_op:
#             style_des = ''
#             style_des_add = []
#             if self.style_name in ['新样式1', '新样式2']:
#                 style_des = '创建并应用“'+self.style_name+'”'
#             else:
#                 style_des = '应用样式“'+self.style_name+'”'

#             font_setting_list=[]
#             if self.style_font_name_chinese !='': font_setting_list.append('中文'+self.style_font_name_chinese)
#             if self.style_font_name_ascii !='':  font_setting_list.append('西文'+self.style_font_name_ascii)
#             if self.style_font_size !='': font_setting_list.append('字号'+self.style_font_size)
#             if self.style_font_color !='': font_setting_list.append(self.style_font_color)
#             if self.style_font_bold==True: font_setting_list.append('粗体')
#             if self.style_font_italic==True: font_setting_list.append('斜体')
#             if self.style_font_underline !='': font_setting_list.append(self.style_font_underline)
#             if len(font_setting_list)>0:
#                 style_des_add.append('其字体设置成'+'、'.join(font_setting_list))

#             para_setting_list=[]
#             if self.style_para_alignment !='': para_setting_list.append(self.style_para_alignment)
#             if self.style_para_left_indent !='': para_setting_list.append('左缩进'+self.style_para_left_indent+'磅')
#             if self.style_para_right_indent !='': para_setting_list.append('右缩进'+self.style_para_right_indent+'磅')
#             if self.style_para_first_line_indent !='' and \
#                self.style_para_first_line_indent_size !='': 
#                 para_setting_list.append(self.style_para_first_line_indent+self.style_para_first_line_indent_size+'磅')
#             if self.style_para_space_before !='': para_setting_list.append('段前'+self.style_para_space_before+'磅')
#             if self.style_para_space_after !='':  para_setting_list.append('段后'+self.style_para_space_before+'磅')
#             if self.style_para_line_spacing_rule !='': 
#                 if self.style_para_line_spacing_rule in ('单倍行距','双倍行距','1.5倍行距'):
#                     para_setting_list.append(self.style_para_line_spacing_rule)
#                 else:
#                     para_setting_list.append(self.style_para_line_spacing+'倍行距')
#             # if self.style_para_firstchardropcap !='' and self.style_para_firstchardropcaplines !='': 
#                 # para_setting_list.append('首字'+self.style_para_firstchardropcap+self.style_para_firstchardropcaplines+'磅')
#             if self.style_page_break_before==True: para_setting_list.append('段前分页')
#             if self.style_keep_with_next==True: para_setting_list.append('与下段同页')
#             if self.style_keep_together==True: para_setting_list.append('段中不分页')
#             if self.style_widow_control==True: para_setting_list.append('孤行控制')

#             if len(para_setting_list)>0:
#                 style_des_add.append('其段落格式设置成'+'、'.join(para_setting_list))
#             return style_des+'('+'，'.join(style_des_add)+')'
#         else:
#             return ''
#     style_description.short_description = '样式设置描述'

#     def image_description(self):
#         if self.image_op:
#             image_desc = '在段落【'+self.para_text_simple()+'】后以'+self.image_position_style+'格式插入图片"'+self.upload_image_file.name.split('/')[-1]+'"'
#             if self.image_width !='':
#                 image_desc = image_desc + '、宽度'+self.image_width+'厘米'
#             if self.image_height !='': 
#                 image_desc = image_desc + '、高度'+self.image_height+'厘米'
#             return image_desc
#         else:
#             return ''
#     image_description.short_description = '图片插入描述'

#     def table_description(self):
#         if self.table_op:
#             table_desc = '将内容【'+self.para_text_simple()+'】生成表格'
#             if self.table_alignment!='':
#                 table_desc = table_desc +'， 表格'+ [x[1] for x in TABLE_ALIGNMENT_CHOICES if x[0]==self.table_alignment][0]
#             if self.table_style!='':
#                 table_desc = table_desc +'， 应用'+ [x[1] for x in TABLE_STYLE_CHOICES if x[0]==self.table_style][0] +'表格样式'
#             if self.table_autofit:
#                 table_desc = table_desc +'， 宽度自动调整' 
#             return table_desc
#         else:
#             return ''
#     table_description.short_description = '表格插入描述'

#     def clean(self):
#         if not (self.char_edit_op or self.font_op or self.paraformat_op or self.style_op or self.image_op or self.table_op):
#             raise ValidationError(_('至少选择一个操作考查'))

#         if self.table_op and (self.char_edit_op or self.font_op or self.paraformat_op or self.style_op or self.image_op):
#             raise ValidationError(_('插入表格不能和其他操作同时选择'))

#         if (self.font_op or self.paraformat_op) and self.style_op:
#             raise ValidationError(_('样式和单独的字体和段落格式不应同时设置'))

#         # check word file
#         # print('file:', self.word_question.upload_docx.upload.path)
#         try:
#             document = Document(self.word_question.upload_docx.upload.path)
#         except:
#             raise ValidationError({'word_question':_('文件异常')})

#         all_paras  = document.paragraphs
#         para_text = self.para_text.strip()
#         matched = 0
#         for para in para_text.split('\n'):
#             for i in range(len(all_paras)):
#                 if all_paras[i].text.strip()==para.strip():
#                     matched = matched + 1
#         if matched != len(para_text.split('\n')):
#             raise ValidationError({'para_text':_('未找到文章中对应内容，请检查是否输入正确内容')})

#         error_dict = {}

#         if self.char_edit_op and (self.char_edit_origin==''):
#             error_dict['char_edit_origin'] = _('不能为空')
#         if self.char_edit_op and (self.char_edit_replace==''):
#             error_dict['char_edit_replace'] = _('不能为空')

#         if self.font_op and \
#         (self.font_name_ascii=='' and 
#         self.font_name_chinese=='' and 
#         self.font_size=='' and 
#         self.font_underline=='' and 
#         self.font_color=='' and
#         self.font_bold==False and
#         self.font_italic==False):
#             error_dict['font_name_chinese'] = _('至少设定一个字体相关设置')
#             error_dict['font_name_ascii'] = _('')
#             error_dict['font_size'] = _('')
#             error_dict['font_underline'] = _('')
#             error_dict['font_color'] = _('')

#         if self.paraformat_op and \
#         (self.para_alignment=='' and 
#         self.para_left_indent=='' and 
#         self.para_right_indent=='' and 
#         self.para_first_line_indent=='' and 
#         self.para_first_line_indent_size=='' and 
#         self.para_space_before=='' and 
#         self.para_space_after=='' and 
#         self.para_line_spacing_rule=='' and 
#         self.para_line_spacing=='' and 
#         # self.para_firstchardropcap=='' and 
#         # self.para_firstchardropcaplines=='' and 
#         self.page_break_before==False and
#         self.keep_with_next==False and
#         self.keep_together==False and
#         self.widow_control==False):
#             error_dict['para_alignment'] = _('至少选择一个段落格式相关设置')
#             error_dict['para_left_indent'] = _('')
#             error_dict['para_right_indent'] = _('')
#             error_dict['para_first_line_indent'] = _('')
#             error_dict['para_first_line_indent_size'] = _('')
#             error_dict['para_space_before'] = _('')
#             error_dict['para_space_after'] = _('')
#             error_dict['para_line_spacing_rule'] = _('')
#             error_dict['para_line_spacing'] = _('')
#             # error_dict['para_firstchardropcap'] = _('')
#             # error_dict['para_firstchardropcaplines'] = _('')

#         if self.paraformat_op and (self.para_first_line_indent=='' and self.para_first_line_indent_size!=''):
#             error_dict['para_first_line_indent'] = _('不能为空')
#             error_dict['para_first_line_indent_size'] = _('*')
#         if self.paraformat_op and (self.para_first_line_indent !='' and self.para_first_line_indent_size==''):
#             error_dict['para_first_line_indent'] = _('*')
#             error_dict['para_first_line_indent_size'] = _('不能为空')

#         if self.paraformat_op and (self.para_line_spacing_rule=='' and self.para_line_spacing!=''):
#             error_dict['para_line_spacing_rule'] = _('不能为空')
#             error_dict['para_line_spacing'] = _('*')
#         if self.paraformat_op and (self.para_line_spacing_rule =='MULTIPLE (5)' and self.para_line_spacing==''):
#             error_dict['para_line_spacing_rule'] = _('*')
#             error_dict['para_line_spacing'] = _('不能为空')

#         # if self.paraformat_op and (self.para_firstchardropcap=='' and self.para_firstchardropcaplines!=''):
#         #     error_dict['para_firstchardropcap'] = _('不能为空')
#         #     error_dict['para_firstchardropcaplines'] = _('*')
#         # if self.paraformat_op and (self.para_firstchardropcap !='' and self.para_firstchardropcaplines==''):
#         #     error_dict['para_firstchardropcap'] = _('*')
#         #     error_dict['para_firstchardropcaplines'] = _('不能为空')


#         if self.style_op and self.style_name=='':
#             error_dict['style_name'] = _('内容不能为空')

#         if self.style_op and self.style_name in ['新样式1', '新样式2'] and \
#                (self.style_font_name_ascii=='' and 
#                 self.style_font_name_chinese=='' and 
#                 self.style_font_size=='' and 
#                 self.style_font_underline=='' and 
#                 self.style_font_color=='' and
#                 self.style_font_bold==False and
#                 self.style_font_italic==False and
#                 self.style_para_alignment=='' and 
#                 self.style_para_left_indent=='' and 
#                 self.style_para_right_indent=='' and 
#                 self.style_para_first_line_indent=='' and 
#                 self.style_para_first_line_indent_size=='' and 
#                 self.style_para_space_before=='' and 
#                 self.style_para_space_after=='' and 
#                 self.style_para_line_spacing_rule=='' and 
#                 self.style_para_line_spacing=='' and 
#                 self.style_page_break_before==False and
#                 self.style_keep_with_next==False and
#                 self.style_keep_together==False and
#                 self.style_widow_control==False
#                 ):
#             error_dict['style_name'] = _('至少为新样式设定一个具体设置')
#             error_dict['style_font_name_chinese'] = _('')
#             error_dict['style_font_name_ascii'] = _('')
#             error_dict['style_font_size'] = _('')
#             error_dict['style_font_underline'] = _('')
#             error_dict['style_font_color'] = _('')
#             error_dict['style_para_alignment'] = _('')
#             error_dict['style_para_left_indent'] = _('')
#             error_dict['style_para_right_indent'] = _('')
#             error_dict['style_para_first_line_indent'] = _('')
#             error_dict['style_para_first_line_indent_size'] = _('')
#             error_dict['style_para_space_before'] = _('')
#             error_dict['style_para_space_after'] = _('')
#             error_dict['style_para_line_spacing_rule'] = _('')
#             error_dict['style_para_line_spacing'] = _('')

#         if self.style_op and \
#         (self.style_para_first_line_indent=='' and 
#          self.style_para_first_line_indent_size!=''):
#             error_dict['style_para_first_line_indent'] = _('不能为空')
#             error_dict['style_para_first_line_indent_size'] = _('*')

#         if self.style_op and \
#             (self.style_para_first_line_indent !='' and 
#              self.style_para_first_line_indent_size==''):
#             error_dict['style_para_first_line_indent'] = _('*')
#             error_dict['style_para_first_line_indent_size'] = _('不能为空')

#         if self.style_op and (self.style_para_line_spacing_rule=='' and self.style_para_line_spacing!=''):
#             error_dict['style_para_line_spacing'] = _('*')
#             error_dict['style_para_line_spacing_rule'] = _('不能为空')
#         if self.style_op and (self.style_para_line_spacing_rule !='' and self.style_para_line_spacing==''):
#             error_dict['style_para_line_spacing'] = _('不能为空')
#             error_dict['style_para_line_spacing_rule'] = _('*')

#         # print('self.upload_image_file',  self.upload_image_file.name)
#         if self.image_op and self.upload_image_file.name=='':
#             error_dict['upload_image_file'] = _('必须上传图片(jpg,bmp,png)')
        
#         if len(error_dict)>0:
#             raise ValidationError(error_dict)


#     ############## 所属大题
#     word_question = models.ForeignKey(
#         WordQuestion,
#         on_delete=models.CASCADE, blank=True, default='',
#         verbose_name='word操作大题'
#     )

#     para_text = models.TextField('要考查的段落内容')

#     ############## 文字查找替换
#     char_edit_op = models.BooleanField('考查文字查找替换？', default=False)
#     char_edit_origin = models.CharField('原词', max_length=200, blank=True, default='')
#     char_edit_replace = models.CharField('替换为', max_length=200, blank=True, default='')

#     ############## 字体设置
#     font_op = models.BooleanField('考查字体设置？', default=False)
#     font_name_ascii = models.CharField('英文字体', choices=FONT_NAME_ASCII_CHOICES, max_length=200, blank=True, default='') 
#     font_name_chinese = models.CharField('中文字体', choices=FONT_NAME_CHINESE_CHOICES, max_length=200, blank=True, default='')
#     font_size = models.CharField('字号(磅)', choices=FONT_SIZE_CHOICES, max_length=200, blank=True, default='')
#     font_bold = models.BooleanField('粗体', blank=True, default='')
#     font_italic = models.BooleanField('斜体', blank=True, default='')
#     font_underline = models.CharField('下划线', choices=FONT_UNDERLINE_CHOICES, max_length=200, blank=True, default='')
#     font_color = models.CharField('字色(标准色)', choices=FONT_COLOR_CHOICES, max_length=200, blank=True, default='')

#     ############## 段落格式设置
#     paraformat_op = models.BooleanField('考查段落格式设置？', default=False)
#     para_alignment = models.CharField('段落对齐', choices=PARA_ALIGNMENT_CHOICES, max_length=200, blank=True, default='')
#     para_left_indent = models.CharField('左侧缩进(磅)', choices=INDENT_NUM_CHOICES, max_length=200, blank=True, default='')
#     para_right_indent = models.CharField('右侧缩进(磅)', choices=INDENT_NUM_CHOICES, max_length=200, blank=True, default='')
#     para_first_line_indent = models.CharField('首行缩进', choices=PARA_FIRST_LINE_INDENT_CHOICES, max_length=200, blank=True, default='')
#     para_first_line_indent_size = models.CharField('首行缩进距离(磅)', choices=INDENT_NUM_CHOICES, max_length=200, blank=True, default='')

#     para_space_before = models.CharField('段前间距(磅)',  choices=INDENT_NUM_CHOICES, max_length=200, blank=True, default='')
#     para_space_after = models.CharField('段后间距(磅)',  choices=INDENT_NUM_CHOICES, max_length=200, blank=True, default='')
#     para_line_spacing_rule = models.CharField('行距规则', choices=PARA_LINE_SPACING_RULE_CHOICES, max_length=200, blank=True, default='')
#     para_line_spacing = models.CharField('行距(行)',  choices=LINE_NUM_CHOICES, max_length=200, blank=True, default='')

#     # para_firstchardropcap = models.CharField('首字下沉', choices=PARA_FIRSTCHARDROPCAP_CHOICES, max_length=200, blank=True, default='')
#     # para_firstchardropcaplines = models.CharField('下沉(行)',  choices=LINE_NUM_CHOICES, max_length=200, blank=True, default='')

#     page_break_before = models.BooleanField('段前分页', default=False)
#     keep_with_next = models.BooleanField('与下段同页', default=False)
#     keep_together = models.BooleanField('段中不分页', default=False)
#     widow_control = models.BooleanField('孤行控制', default=False)

#     ############## 样式设置
#     style_op = models.BooleanField('考查样式设置？', default=False)
#     style_name = models.CharField('样式名称', choices=STYLE_NAME_CHOICES, max_length=200, blank=True, default='') 
#     style_font_name_ascii = models.CharField('英文字体', choices=FONT_NAME_ASCII_CHOICES, max_length=200, blank=True, default='') 
#     style_font_name_chinese = models.CharField('中文字体', choices=FONT_NAME_CHINESE_CHOICES, max_length=200, blank=True, default='')
#     style_font_size = models.CharField('字号', choices=FONT_SIZE_CHOICES, max_length=200, blank=True, default='')
#     style_font_bold = models.BooleanField('粗体', default=False)
#     style_font_italic = models.BooleanField('斜体', default=False)
#     style_font_underline = models.CharField('下划线', choices=FONT_UNDERLINE_CHOICES, max_length=200, blank=True, default='')
#     style_font_color = models.CharField('字色(标准色)', choices=FONT_COLOR_CHOICES, max_length=200, blank=True, default='')

#     style_para_alignment = models.CharField('段落对齐', choices=PARA_ALIGNMENT_CHOICES, max_length=200, blank=True, default='')
#     style_para_left_indent = models.CharField('左侧缩进(磅)',  choices=INDENT_NUM_CHOICES, max_length=200, blank=True, default='')
#     style_para_right_indent = models.CharField('右侧缩进(磅)',  choices=INDENT_NUM_CHOICES, max_length=200, blank=True, default='')
#     style_para_first_line_indent = models.CharField('首行缩进', choices=PARA_FIRST_LINE_INDENT_CHOICES, max_length=200, blank=True, default='')
#     style_para_first_line_indent_size = models.CharField('首行缩进距离(磅)',  choices=INDENT_NUM_CHOICES, max_length=200, blank=True, default='')

#     style_para_space_before = models.CharField('段前间距(磅)',   choices=INDENT_NUM_CHOICES, max_length=200, blank=True, default='')
#     style_para_space_after = models.CharField('段后间距(磅)',   choices=INDENT_NUM_CHOICES, max_length=200, blank=True, default='')
#     style_para_line_spacing_rule = models.CharField('行距规则', choices=PARA_LINE_SPACING_RULE_CHOICES, max_length=200, blank=True, default='')
#     style_para_line_spacing = models.CharField('行距(行)',   choices=LINE_NUM_CHOICES, max_length=200, blank=True, default='')

#     style_page_break_before = models.BooleanField('段前分页', default=False)
#     style_keep_with_next = models.BooleanField('与下段同页', default=False)
#     style_keep_together = models.BooleanField('段中不分页', default=False)
#     style_widow_control = models.BooleanField('孤行控制', default=False)

#     ############## 图片插入
#     image_op = models.BooleanField('考查图片插入？', default=False)
#     image_position_style = models.CharField('位置类型', choices=IMAGE_POSITION_STYLE_CHOICES, max_length=20, default='嵌入文本行中')
#     image_width = models.CharField('图像宽(厘米)',   choices=INDENT_NUM_CHOICES, max_length=200, blank=True, default='')
#     image_height = models.CharField('图像高(厘米)',   choices=INDENT_NUM_CHOICES, max_length=200, blank=True, default='')
#     upload_image_file = models.FileField(upload_to='uploads_image/', null=True, blank=True, 
#     validators=[validate_image], verbose_name='上传图片')

#     ############## 表格插入
#     table_op = models.BooleanField('考查表格插入？', default=False)
#     table_alignment = models.CharField('对齐方式', choices=TABLE_ALIGNMENT_CHOICES, max_length=10, default='')
#     table_autofit = models.BooleanField('表格宽度自动调整？', default=False)
#     table_style = models.CharField('表格样式', choices=TABLE_STYLE_CHOICES, max_length=30, default='')

